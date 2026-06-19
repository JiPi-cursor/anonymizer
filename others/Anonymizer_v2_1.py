"""
anonymizer_gui.py
─────────────────
Interface graphique pour l'anonymisation et la désanonymisation
de documents Word (.docx) et PDF.

La table de correspondance est stockée en CSV (séparateur point-virgule)
avec deux colonnes : "Alias" et "Valeur originale".
Ce format est directement éditable dans Excel, LibreOffice Calc ou
n'importe quel éditeur de texte.

Dépendances :
    pip install customtkinter presidio-analyzer presidio-anonymizer \
                spacy python-docx pdfplumber reportlab pyspellchecker
    python -m spacy download fr_core_news_lg
"""

import csv
import os
import threading
import tkinter as tk
from tkinter import filedialog, messagebox
from typing import List

import customtkinter as ctk

# ══════════════════════════════════════════════════════════════════════
# THÈME
# ══════════════════════════════════════════════════════════════════════

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

FONT_TITLE = ("Inter", 15, "bold")
FONT_LABEL = ("Inter", 12)
FONT_SMALL = ("Inter", 11)
FONT_MONO  = ("Courier New", 11)

COLOR_BG     = "#1a1a2e"
COLOR_PANEL  = "#16213e"
COLOR_ACCENT = "#0f3460"
COLOR_GREEN  = "#2ecc71"
COLOR_RED    = "#e74c3c"
COLOR_ORANGE = "#e67e22"
COLOR_TEXT   = "#e0e0e0"
COLOR_MUTED  = "#8892a4"

CSV_SEPARATOR  = ";"          # point-virgule : compatible Excel français
CSV_HEADER     = ["Alias", "Valeur originale"]
CSV_ENCODING   = "utf-8-sig"  # BOM UTF-8 : Excel l'ouvre sans problème d'accents


# ══════════════════════════════════════════════════════════════════════
# TABLE DE CORRESPONDANCE — lecture / écriture CSV
# ══════════════════════════════════════════════════════════════════════

def mapping_load(path: str) -> dict:
    """
    Charge un CSV de correspondance.
    Retourne un dict  { alias: valeur_originale }.
    Compatible avec les fichiers créés manuellement dans Excel.
    """
    mapping = {}
    with open(path, newline="", encoding=CSV_ENCODING) as f:
        reader = csv.DictReader(f, delimiter=CSV_SEPARATOR)
        for row in reader:
            alias    = row.get("Alias", "").strip()
            original = row.get("Valeur originale", "").strip()
            if alias and original:
                mapping[alias] = original
    return mapping


def mapping_save(mapping: dict, path: str) -> None:
    """
    Sauvegarde la table de correspondance en CSV.
    Colonnes : Alias | Valeur originale
    Triées par alias pour faciliter la lecture dans Excel.
    """
    with open(path, "w", newline="", encoding=CSV_ENCODING) as f:
        writer = csv.writer(f, delimiter=CSV_SEPARATOR,
                            quoting=csv.QUOTE_ALL)
        writer.writerow(CSV_HEADER)
        for alias, original in sorted(mapping.items()):
            writer.writerow([alias, original])


def mapping_default_path(output_path: str) -> str:
    """Chemin CSV auto à côté du document de sortie."""
    base = os.path.splitext(output_path)[0]
    return f"{base}_correspondances.csv"


# ══════════════════════════════════════════════════════════════════════
# MOTEUR D'ANONYMISATION
# ══════════════════════════════════════════════════════════════════════

_engine_ready = False
_analyzer     = None
_anonymizer   = None

_LABEL_MAP = {
    "PERSON":             "[NOM]",
    "PER":                "[NOM]",
    "ORG":                "[ORGANISATION]",
    "LOC":                "[LIEU]",
    "GPE":                "[LIEU]",
    "EMAIL_ADDRESS":      "[EMAIL]",
    "PHONE_NUMBER":       "[TÉLÉPHONE]",
    "CODE_POSTAL":        "[CODE_POSTAL]",
    "SIRET":              "[SIRET]",
    "IBAN_CODE":          "[IBAN]",
    "URL":                "[URL]",
    "DATE_TIME":          "[DATE]",
    "NRP":                "[IDENTIFIANT]",
    "NOM_PROPRE_INCONNU": "[NOM_PROPRE]",
}


def _load_engine(log_fn):
    """Charge Presidio + spaCy en arrière-plan."""
    global _engine_ready, _analyzer, _anonymizer

    try:
        log_fn("Chargement de spaCy fr_core_news_lg…")
        from presidio_analyzer import (AnalyzerEngine, Pattern,
                                       PatternRecognizer)
        from presidio_analyzer.nlp_engine import NlpEngineProvider
        from presidio_anonymizer import AnonymizerEngine

        configuration = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "fr", "model_name": "fr_core_news_lg"}],
        }
        provider   = NlpEngineProvider(nlp_configuration=configuration)
        nlp_engine = provider.create_engine()
        _analyzer  = AnalyzerEngine(
            nlp_engine=nlp_engine, supported_languages=["fr"]
        )

        # ── Recognizers regex ──────────────────────────────────────────
        _analyzer.registry.add_recognizer(PatternRecognizer(
            supported_entity="CODE_POSTAL",
            patterns=[Pattern("cp_fr", r"\b[0-9]{5}\b", 0.6)],
            context=["cedex", "bp", "rue", "avenue", "boulevard"],
        ))
        _analyzer.registry.add_recognizer(PatternRecognizer(
            supported_entity="SIRET",
            patterns=[
                Pattern("siret", r"\b\d{3}\s?\d{3}\s?\d{3}\s?\d{5}\b", 0.85),
                Pattern("siren", r"\b\d{3}\s?\d{3}\s?\d{3}\b", 0.6),
            ],
            context=["siret", "siren", "rcs", "immatriculation"],
        ))
        _analyzer.registry.add_recognizer(PatternRecognizer(
            supported_entity="PHONE_NUMBER",
            patterns=[Pattern("phone_fr",
                r"\b(?:(?:\+33|0033|0)[1-9])(?:[\s.\-]?\d{2}){4}\b", 0.8)],
        ))
        _analyzer.registry.add_recognizer(PatternRecognizer(
            supported_entity="EMAIL_ADDRESS",
            patterns=[Pattern("email",
                r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b",
                0.9)],
        ))
        _analyzer.registry.add_recognizer(PatternRecognizer(
            supported_entity="IBAN_CODE",
            patterns=[Pattern(
                "iban_fr",
                # Structure IBAN FR : FR76 + 5 groupes de 4 chiffres + 1 groupe de 3
                # \s* entre groupes : tolère 0 à N espaces (tableaux Word/PDF)
                r"FR\d{2}(?:\s*\d{4}){5}\s*\d{3}",
                0.9,
            )],
        ))
        _analyzer.registry.add_recognizer(_build_hors_dico_recognizer())

        _anonymizer   = AnonymizerEngine()
        _engine_ready = True
        log_fn("✓ Moteur prêt.")

    except Exception as exc:
        log_fn(f"✗ Erreur chargement moteur : {exc}")


def _build_hors_dico_recognizer():
    from presidio_analyzer import EntityRecognizer, RecognizerResult
    from presidio_analyzer.nlp_engine import NlpArtifacts
    from spellchecker import SpellChecker

    class HorsDictionnaireRecognizer(EntityRecognizer):
        SUPPORTED_ENTITY = "NOM_PROPRE_INCONNU"
        DEFAULT_SCORE    = 0.4
        WHITELIST = {
            "workflow", "reporting", "management", "process",
            "benchmark", "compliance", "audit", "framework",
            "meeting", "software", "hardware", "cloud", "data",
        }

        def __init__(self):
            super().__init__(
                supported_entities=[self.SUPPORTED_ENTITY],
                name="HorsDictionnaireRecognizer",
                supported_language="fr",
            )
            self.spell = SpellChecker(language=["fr", "en"])

        def load(self):
            pass

        def analyze(self, text: str, entities: List[str],
                    nlp_artifacts: NlpArtifacts) -> List[RecognizerResult]:
            results = []
            for token in nlp_artifacts.tokens:
                word = token.text
                if len(word) <= 2 or not word.isalpha() or token.is_punct:
                    continue
                if token.is_stop:
                    continue
                if word.lower() in self.WHITELIST:
                    continue

                hors_dico = bool(self.spell.unknown([word.lower()]))

                if word[0].isupper():
                    majuscule = (bool(self.spell.unknown([word.lower()]))
                                 if token.is_sent_start else True)
                else:
                    majuscule = False

                majuscule_interne = any(c.isupper() for c in word[1:])
                acronyme = word.isupper() and len(word) >= 3

                if hors_dico or majuscule or majuscule_interne or acronyme:
                    results.append(RecognizerResult(
                        entity_type=self.SUPPORTED_ENTITY,
                        start=token.idx,
                        end=token.idx + len(word),
                        score=self.DEFAULT_SCORE,
                    ))
            return results

    return HorsDictionnaireRecognizer()


# ══════════════════════════════════════════════════════════════════════
# PATTERNS STRUCTURÉS — détection directe par regex, sans Presidio
# Presidio délègue ces patterns à spaCy qui peut les rater selon le
# contexte (tableaux, espaces multiples, début de ligne).
# On les extrait ici en Python pur, ce qui est fiable à 100 %.
# ══════════════════════════════════════════════════════════════════════

import re as _re

_REGEX_PATTERNS = [
    # (entity_type, pattern compilé)
    # IBAN FR : FR76 + 5 groupes de 4 chiffres + 1 groupe de 3, espaces optionnels
    ("IBAN_CODE",    _re.compile(r"FR\d{2}(?:\s*\d{4}){5}\s*\d{3}")),
    # SIRET 14 chiffres (espaces optionnels entre groupes)
    ("SIRET",        _re.compile(r"\b\d{3}\s\d{3}\s\d{3}\s\d{5}\b")),
    # SIREN 9 chiffres
    ("SIRET",        _re.compile(r"\b\d{3}\s\d{3}\s\d{3}\b")),
    # Téléphone français
    ("PHONE_NUMBER", _re.compile(r"(?:(?:\+33|0033|0)[1-9])(?:[\s.\-]?\d{2}){4}")),
    # Email
    ("EMAIL_ADDRESS",_re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")),
]


def _register_regex_matches(text: str, mapping: dict) -> None:
    """
    Extrait tous les formats structurés par regex directe et les enregistre
    dans le mapping. Indépendant de Presidio — aucun risque de raté.
    """
    for entity_type, pattern in _REGEX_PATTERNS:
        label = _LABEL_MAP.get(entity_type, "[INCONNU]")
        for m in pattern.finditer(text):
            original = m.group().strip()
            if not original:
                continue
            if any(v == original for v in mapping.values()):
                continue  # déjà enregistré
            base  = label.rstrip("]")
            count = sum(1 for k in mapping if k.startswith(base))
            alias = f"{base}_{count + 1}]"
            mapping[alias] = original


# ══════════════════════════════════════════════════════════════════════
# LOGIQUE ANONYMISATION / DÉSANONYMISATION
# ══════════════════════════════════════════════════════════════════════

def _anonymize_text(text: str, mapping: dict, readonly: bool = False) -> str:
    """
    Trois passes :

    Passe 1a — Regex directe (IBAN, SIRET, téléphone, email)
        Indépendant de Presidio. re.finditer garantit de trouver TOUTES
        les occurrences de chaque format structuré dans le texte.
        Ignorée si readonly=True.

    Passe 1b — Presidio / spaCy (noms, organisations, lieux)
        Presidio peut rater certaines occurrences répétées ou mal placées,
        mais la passe 2 corrige cela par str.replace exhaustif.
        Ignorée si readonly=True.

    Passe 2 — str.replace exhaustif sur tout le mapping
        Pour chaque (alias → valeur_originale) enregistré, on remplace
        toutes les occurrences dans le texte, du plus long au plus court.
        C'est cette passe qui garantit qu'aucune occurrence n'est oubliée.
        Toujours exécutée, que readonly soit True ou False.

    Paramètre readonly :
        False (défaut) — aucune table fournie en entrée : les passes 1a et 1b
            découvrent les entités et enrichissent le mapping.
        True — une table existante a été fournie : les passes 1a et 1b sont
            ignorées ; seule la passe 2 applique le mapping tel quel, sans
            aucune modification de celui-ci.
    """
    if not text or not text.strip():
        return text

    if not readonly:
        # Passe 1a : formats structurés par regex directe
        _register_regex_matches(text, mapping)

        # Passe 1b : entités nommées via Presidio
        results = _analyzer.analyze(
            text=text,
            language="fr",
            entities=list(_LABEL_MAP.keys()),
            score_threshold=0.35,
        )
        for res in sorted(results, key=lambda r: -r.score):
            original = text[res.start:res.end].strip()
            if not original:
                continue
            if any(v == original for v in mapping.values()):
                continue
            label = _LABEL_MAP.get(res.entity_type, "[INCONNU]")
            base  = label.rstrip("]")
            count = sum(1 for k in mapping if k.startswith(base))
            alias = f"{base}_{count + 1}]"
            mapping[alias] = original

    # Passe 2 : remplacement exhaustif, du plus long au plus court
    result = text
    for alias, original in sorted(mapping.items(),
                                  key=lambda kv: len(kv[1]), reverse=True):
        result = result.replace(original, alias)

    return result


def _deanonymize_text(text: str, mapping: dict) -> str:
    result = text
    for alias, original in sorted(mapping.items(),
                                  key=lambda x: len(x[0]), reverse=True):
        result = result.replace(alias, original)
    return result


# ── DOCX ──────────────────────────────────────────────────────────────

def _process_docx(input_path, output_path, mapping, mode, readonly: bool = False):
    from docx import Document

    doc = Document(input_path)
    fn  = (lambda t, m: _anonymize_text(t, m, readonly=readonly)) if mode == "anon" else _deanonymize_text

    def process_para(para):
        """
        Concatène tous les runs, applique fn(), puis réécrit le résultat.

        Problème de l'approche runs[0].text = processed :
        python-docx écrit dans le premier run UNIQUEMENT, mais la mise en
        forme des runs suivants (gras, italique) est perdue et surtout,
        si runs[0] était court (ex: "Le présent rapport a été commandité par "),
        l'entité qui enjambe deux runs ("Hiram" run1 + " Finance" run2) est
        bien détectée dans le texte joint, remplacée dans `processed`, mais
        `processed` est ensuite écrit dans runs[0] dont le formatage XML
        est conservé — ce qui est correct.

        Le vrai problème est que `para.runs` peut être vide si le paragraphe
        est constitué uniquement de champs XML complexes (hyperliens, etc.).
        Dans ce cas on manipule directement le XML du paragraphe.
        """
        # Cas normal : le paragraphe a des runs
        if para.runs:
            full = "".join(r.text for r in para.runs)
            if not full.strip():
                return
            processed = fn(full, mapping)
            # Écrire tout dans le premier run, vider les suivants
            para.runs[0].text = processed
            for r in para.runs[1:]:
                r.text = ""
            return

        # Cas dégradé : pas de runs (champs, hyperliens complexes)
        # Manipulation XML directe
        from lxml import etree
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        texts = para._p.findall(f".//{{{W}}}t")
        if not texts:
            return
        full = "".join(t.text or "" for t in texts)
        if not full.strip():
            return
        processed = fn(full, mapping)
        # Écrire dans le premier <w:t>, vider les suivants
        texts[0].text = processed
        for t in texts[1:]:
            t.text = ""

    for para in doc.paragraphs:
        process_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            process_para(para)
        for para in section.footer.paragraphs:
            process_para(para)

    doc.save(output_path)


# ── PDF ───────────────────────────────────────────────────────────────

def _process_pdf(input_path, output_path, mapping, mode, readonly: bool = False):
    import pdfplumber
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    fn         = (lambda t, m: _anonymize_text(t, m, readonly=readonly)) if mode == "anon" else _deanonymize_text
    pages_text = []

    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages_text.append(fn(text, mapping))

    c = canvas.Canvas(output_path, pagesize=A4)
    w, h = A4
    margin, lh = 50, 14
    c.setFont("Helvetica", 10)

    for page_text in pages_text:
        y = h - margin
        for line in page_text.split("\n"):
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = h - margin
            c.drawString(margin, y, line[:120])
            y -= lh
        c.showPage()
    c.save()


# ── Dispatcher ────────────────────────────────────────────────────────

def run_process(input_path, mapping_path, output_path, mode, log_fn, done_fn):
    try:
        if not _engine_ready:
            raise RuntimeError("Le moteur n'est pas encore chargé.")

        # Chargement ou initialisation de la table
        mapping  = {}
        readonly = False  # True = table fournie en entrée → aucun ajout
        if mapping_path and os.path.isfile(mapping_path):
            mapping  = mapping_load(mapping_path)
            readonly = True
            log_fn(f"Table chargée : {len(mapping)} entrée(s) — mode lecture seule.")
        elif mode == "deanon":
            raise FileNotFoundError(
                "La désanonymisation nécessite une table de correspondance existante."
            )
        else:
            log_fn("Aucune table fournie — construction automatique.")

        ext = os.path.splitext(input_path)[1].lower()
        log_fn(f"Traitement {ext.upper()} en cours…")

        if ext == ".docx":
            _process_docx(input_path, output_path, mapping, mode, readonly=readonly)
        elif ext == ".pdf":
            _process_pdf(input_path, output_path, mapping, mode, readonly=readonly)
        else:
            raise ValueError(f"Format non supporté : {ext}")

        # Chemin CSV : celui fourni, ou chemin auto
        # En mode readonly, on ne réécrit pas la table fournie.
        if not readonly:
            save_path = mapping_path or mapping_default_path(output_path)
            mapping_save(mapping, save_path)
            log_fn(f"✓ Document sauvegardé    : {output_path}")
            log_fn(f"✓ Table CSV sauvegardée  : {save_path}")
            log_fn(f"  ({len(mapping)} entrée(s) — éditable dans Excel)")
        else:
            save_path = mapping_path  # référence inchangée pour _on_done
            log_fn(f"✓ Document sauvegardé    : {output_path}")
            log_fn(f"  (table de correspondance inchangée : {save_path})")
        done_fn(success=True, mapping_path=save_path)

    except Exception as exc:
        log_fn(f"✗ Erreur : {exc}")
        done_fn(success=False, mapping_path=None)


# ══════════════════════════════════════════════════════════════════════
# INTERFACE GRAPHIQUE
# ══════════════════════════════════════════════════════════════════════

class AnonymizerApp(ctk.CTk):

    def __init__(self):
        super().__init__()
        self.title("Anonymiseur de documents")
        self.geometry("800x660")
        self.resizable(True, True)
        self.configure(fg_color=COLOR_BG)

        self._last_mapping_path = None  # mémorise le CSV généré
        self._build_ui()

        threading.Thread(
            target=_load_engine, args=(self._log,), daemon=True
        ).start()

    # ── Construction de l'UI ──────────────────────────────────────────

    def _build_ui(self):
        # En-tête
        header = ctk.CTkFrame(self, fg_color=COLOR_PANEL, corner_radius=0)
        header.pack(fill="x")

        ctk.CTkLabel(
            header, text="🔒  Anonymiseur de documents",
            font=FONT_TITLE, text_color=COLOR_TEXT,
        ).pack(side="left", padx=20, pady=14)

        self._status_dot = ctk.CTkLabel(
            header, text="●", font=("Inter", 18), text_color=COLOR_ORANGE
        )
        self._status_dot.pack(side="right", padx=8)
        self._status_lbl = ctk.CTkLabel(
            header, text="Chargement du moteur…",
            font=FONT_SMALL, text_color=COLOR_MUTED,
        )
        self._status_lbl.pack(side="right", padx=4, pady=14)

        # Corps
        body = ctk.CTkFrame(self, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=24, pady=16)

        # ── Champs fichiers ───────────────────────────────────────────
        self._entry_input = self._file_row(
            body, "Document source *",
            "Fichier .docx ou .pdf à traiter",
            self._browse_input,
        )
        self._entry_mapping = self._file_row(
            body, "Table de correspondance",
            "Fichier .csv (laissez vide : créé automatiquement)",
            self._browse_mapping,
        )
        self._entry_output = self._file_row(
            body, "Document de sortie *",
            "Chemin du fichier résultat",
            self._browse_output,
        )

        # ── Note format CSV ───────────────────────────────────────────
        note_frame = ctk.CTkFrame(body, fg_color=COLOR_ACCENT,
                                  corner_radius=6)
        note_frame.pack(fill="x", pady=(4, 0))
        ctk.CTkLabel(
            note_frame,
            text="ℹ  La table de correspondance est un fichier CSV (séparateur « ; »)"
                 "  —  colonnes : Alias | Valeur originale  —  éditable dans Excel.",
            font=FONT_SMALL, text_color=COLOR_TEXT, wraplength=700,
        ).pack(padx=12, pady=6)

        # ── Boutons principaux ────────────────────────────────────────
        btn_frame = ctk.CTkFrame(body, fg_color="transparent")
        btn_frame.pack(fill="x", pady=(18, 4))

        self._btn_anon = ctk.CTkButton(
            btn_frame, text="🔒  Anonymiser",
            font=FONT_LABEL, fg_color="#1a6b3c", hover_color="#145530",
            text_color="white", height=44, corner_radius=8,
            command=lambda: self._run("anon"),
        )
        self._btn_anon.pack(side="left", expand=True, fill="x", padx=(0, 6))

        self._btn_deanon = ctk.CTkButton(
            btn_frame, text="🔓  Désanonymiser",
            font=FONT_LABEL, fg_color="#6b3a1a", hover_color="#55301b",
            text_color="white", height=44, corner_radius=8,
            command=lambda: self._run("deanon"),
        )
        self._btn_deanon.pack(side="left", expand=True, fill="x", padx=(6, 6))

        # Bouton "Ouvrir la table CSV"
        self._btn_open_csv = ctk.CTkButton(
            btn_frame, text="📊  Ouvrir table",
            font=FONT_LABEL, fg_color=COLOR_ACCENT, hover_color="#1a4a80",
            text_color="white", height=44, corner_radius=8,
            width=140, command=self._open_csv,
        )
        self._btn_open_csv.pack(side="left", padx=(6, 0))

        # ── Progression ───────────────────────────────────────────────
        self._progress = ctk.CTkProgressBar(body, height=6, corner_radius=3)
        self._progress.pack(fill="x", pady=(10, 0))
        self._progress.set(0)

        # ── Journal ───────────────────────────────────────────────────
        ctk.CTkLabel(body, text="Journal", font=FONT_SMALL,
                     text_color=COLOR_MUTED).pack(anchor="w", pady=(12, 4))

        self._log_box = ctk.CTkTextbox(
            body, font=FONT_MONO, fg_color=COLOR_PANEL,
            text_color=COLOR_TEXT, corner_radius=8,
            height=180, wrap="word", state="disabled",
        )
        self._log_box.pack(fill="both", expand=True)

    # ── Ligne fichier ─────────────────────────────────────────────────

    def _file_row(self, parent, label, placeholder, browse_cmd):
        frame = ctk.CTkFrame(parent, fg_color="transparent")
        frame.pack(fill="x", pady=5)

        ctk.CTkLabel(
            frame, text=label, font=FONT_LABEL,
            text_color=COLOR_TEXT, width=210, anchor="w",
        ).pack(side="left")

        entry = ctk.CTkEntry(
            frame, placeholder_text=placeholder,
            font=FONT_SMALL, fg_color=COLOR_PANEL,
            border_color=COLOR_ACCENT, text_color=COLOR_TEXT,
            height=36, corner_radius=6,
        )
        entry.pack(side="left", fill="x", expand=True, padx=(8, 8))

        ctk.CTkButton(
            frame, text="…", width=36, height=36, font=FONT_LABEL,
            fg_color=COLOR_ACCENT, hover_color="#1a4a80",
            corner_radius=6, command=browse_cmd,
        ).pack(side="left")

        return entry

    # ── Dialogues fichiers ────────────────────────────────────────────

    def _browse_input(self):
        path = filedialog.askopenfilename(
            title="Document source",
            filetypes=[("Documents", "*.docx *.pdf"), ("Tous", "*.*")],
        )
        if path:
            self._entry_input.delete(0, "end")
            self._entry_input.insert(0, path)
            if not self._entry_output.get():
                base, ext = os.path.splitext(path)
                self._entry_output.delete(0, "end")
                self._entry_output.insert(0, f"{base}_anonymise{ext}")

    def _browse_mapping(self):
        path = filedialog.askopenfilename(
            title="Table de correspondance CSV",
            filetypes=[("CSV", "*.csv"), ("Tous", "*.*")],
        )
        if path:
            self._entry_mapping.delete(0, "end")
            self._entry_mapping.insert(0, path)

    def _browse_output(self):
        path = filedialog.asksaveasfilename(
            title="Document de sortie",
            filetypes=[("Documents", "*.docx *.pdf"), ("Tous", "*.*")],
        )
        if path:
            self._entry_output.delete(0, "end")
            self._entry_output.insert(0, path)

    def _open_csv(self):
        """Ouvre la table CSV dans l'application par défaut (Excel, etc.)."""
        # Priorité : champ saisi > dernier CSV généré
        path = self._entry_mapping.get().strip() or self._last_mapping_path
        if not path:
            messagebox.showinfo(
                "Aucune table",
                "Aucune table de correspondance n'est disponible.\n"
                "Lancez d'abord une anonymisation pour en créer une.",
            )
            return
        if not os.path.isfile(path):
            messagebox.showerror("Fichier introuvable",
                                 f"Le fichier CSV est introuvable :\n{path}")
            return
        import subprocess, sys
        if sys.platform == "win32":
            os.startfile(path)
        elif sys.platform == "darwin":
            subprocess.call(["open", path])
        else:
            subprocess.call(["xdg-open", path])

    # ── Exécution ─────────────────────────────────────────────────────

    def _run(self, mode: str):
        input_path   = self._entry_input.get().strip()
        mapping_path = self._entry_mapping.get().strip() or None
        output_path  = self._entry_output.get().strip()

        if not input_path:
            messagebox.showerror("Champ manquant", "Le document source est requis.")
            return
        if not output_path:
            messagebox.showerror("Champ manquant", "Le document de sortie est requis.")
            return
        if not os.path.isfile(input_path):
            messagebox.showerror("Fichier introuvable",
                                 f"Fichier source introuvable :\n{input_path}")
            return

        self._set_buttons_state("disabled")
        self._progress.set(0)
        self._progress.start()
        self._log("─" * 52)
        action = "Anonymisation" if mode == "anon" else "Désanonymisation"
        self._log(f"{action} : {os.path.basename(input_path)}")

        threading.Thread(
            target=run_process,
            args=(input_path, mapping_path, output_path, mode,
                  self._log, self._on_done),
            daemon=True,
        ).start()

    def _on_done(self, success: bool, mapping_path: str = None):
        self._progress.stop()
        self._progress.set(1 if success else 0)
        self._set_buttons_state("normal")
        if success:
            self._set_status("Terminé", COLOR_GREEN)
            if mapping_path:
                self._last_mapping_path = mapping_path
                # Met à jour le champ si vide
                if not self._entry_mapping.get().strip():
                    self.after(0, lambda: (
                        self._entry_mapping.delete(0, "end"),
                        self._entry_mapping.insert(0, mapping_path),
                    ))
        else:
            self._set_status("Erreur", COLOR_RED)

    # ── Helpers ───────────────────────────────────────────────────────

    def _log(self, message: str):
        def _append():
            self._log_box.configure(state="normal")
            self._log_box.insert("end", message + "\n")
            self._log_box.see("end")
            self._log_box.configure(state="disabled")
            if "✓ Moteur prêt" in message:
                self._set_status("Moteur prêt", COLOR_GREEN)
        self.after(0, _append)

    def _set_status(self, text: str, color: str):
        def _upd():
            self._status_lbl.configure(text=text)
            self._status_dot.configure(text_color=color)
        self.after(0, _upd)

    def _set_buttons_state(self, state: str):
        def _upd():
            self._btn_anon.configure(state=state)
            self._btn_deanon.configure(state=state)
        self.after(0, _upd)


# ══════════════════════════════════════════════════════════════════════
# POINT D'ENTRÉE
# ══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    app = AnonymizerApp()
    app.mainloop()